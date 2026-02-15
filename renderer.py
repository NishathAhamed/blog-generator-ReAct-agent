from __future__ import annotations

import os
from typing import Any, Dict

from docx import Document
from docx.shared import Inches


def render_docx(doc_spec: Dict[str, Any], assets_by_id: Dict[str, dict], out_path: str) -> None:
    doc = Document()

    doc.add_heading(doc_spec.get("title", "Untitled"), level=0)
    subtitle = doc_spec.get("subtitle", "")
    if subtitle:
        doc.add_paragraph(subtitle)

    inserted_images = 0

    for section in doc_spec.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=1)

        for p in section.get("paragraphs", []):
            doc.add_paragraph(str(p))

        for img in section.get("images", []):
            asset_id = img.get("asset_id")
            caption = img.get("caption", "")
            asset = assets_by_id.get(asset_id)
            if not asset:
                continue

            path = asset.get("path")
            if not path or not os.path.exists(path):
                continue

            doc.add_picture(path, width=Inches(6.0))
            inserted_images += 1
            if caption:
                cap_p = doc.add_paragraph(caption)
                if cap_p.runs:
                    cap_p.runs[0].italic = True

    # Fallback: if doc_spec didn't reference images, still embed first 2 assets
    if inserted_images == 0 and assets_by_id:
        doc.add_heading("Images", level=1)
        for asset in list(assets_by_id.values())[:2]:
            path = asset.get("path")
            if path and os.path.exists(path):
                doc.add_picture(path, width=Inches(6.0))
                src = asset.get("source_url", "")
                if src:
                    doc.add_paragraph(src)

    refs = doc_spec.get("references", [])
    if refs:
        doc.add_heading("References", level=1)
        for r in refs:
            doc.add_paragraph(f"{r.get('title','Source')} — {r.get('url','')}")

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
