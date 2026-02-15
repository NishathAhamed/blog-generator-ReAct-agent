# dev-test/dev_probe.py

import sys
from pathlib import Path

# ✅ add project root BEFORE importing local modules
ROOT = Path(__file__).resolve().parents[1]  # dev-test -> project root
sys.path.insert(0, str(ROOT))

import json
import zipfile
from docx import Document
from docx.shared import Inches

from tools import web_search, fetch_image  # ✅ now this will work


OUT = ROOT / "output"
OUT.mkdir(exist_ok=True)

print("ROOT:", ROOT)
print("\n1) web_search...")
search = json.loads(web_search.invoke("Retrieval-Augmented Generation RAG diagram png"))
print("images found:", len(search.get("images", [])))
print("first image url:", (search.get("images") or [{}])[0].get("url"))

print("\n2) fetch_image...")
url = (search.get("images") or [{}])[0].get("url")
if not url:
    raise RuntimeError("No image URLs returned from web_search")

asset = json.loads(fetch_image.invoke(url))
path = Path(asset["path"])
print("downloaded path:", path)
print("exists:", path.exists(), "size:", path.stat().st_size if path.exists() else None)

print("\n3) python-docx embed test...")
docx_path = OUT / "_embed_test.docx"
doc = Document()
doc.add_heading("Embed test", level=0)
doc.add_paragraph("If you can see the image below, embedding works.")
doc.add_picture(str(path), width=Inches(5.5))
doc.save(docx_path)

with zipfile.ZipFile(docx_path) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
print("embedded media:", media)
print("OK ->", docx_path)
